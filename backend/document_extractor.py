"""Document text extraction for ingestion pipeline.

Supports: PDF (pdfplumber), DOCX (python-docx), XLSX (openpyxl),
          MD / TXT / plain text (built-in).

All extractors truncate output to MAX_CHARS (50 000) to stay within
the LLM context budget.
"""
from __future__ import annotations

import base64
import os
import pathlib

MAX_CHARS = 50_000

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".md", ".txt", ".rtf",
                        ".png", ".jpg", ".jpeg", ".webp", ".gif"}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}

_VISION_PROMPT = """\
You are analyzing an image uploaded to a knowledge base.
Describe every component, relationship, data flow, decision, and label visible.
Output structured text (headings, bullet points) suitable for a wiki page.
Include: what the diagram shows, every named component, every arrow/connection and what it means,
any labels, annotations, or decision points. Be exhaustive — nothing visible should be omitted.\
"""


def _get_anthropic_client():
    import anthropic
    return anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))


class UnsupportedFileType(ValueError):
    pass


def extract_document(file_path: str) -> dict:
    """Dispatch to the right extractor based on file extension."""
    ext = pathlib.Path(file_path).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(f"Unsupported file type: {ext!r}")
    if ext == ".pdf":
        return extract_pdf(file_path)
    if ext == ".docx":
        return extract_docx(file_path)
    if ext == ".xlsx":
        return extract_xlsx(file_path)
    if ext in IMAGE_EXTENSIONS:
        return extract_image(file_path)
    # .md, .txt, .rtf — plain text
    return extract_text_file(file_path)


def extract_pdf(file_path: str) -> dict:
    import pdfplumber

    pages_text: list[str] = []
    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            t = page.extract_text() or ""
            pages_text.append(t)

    full = "\n".join(pages_text)
    truncated = len(full) > MAX_CHARS
    return {
        "text": full[:MAX_CHARS],
        "page_count": page_count,
        "char_count": len(full),
        "truncated": truncated,
    }


def extract_docx(file_path: str) -> dict:
    from docx import Document  # python-docx

    doc = Document(file_path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    has_tables = bool(doc.tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)

    full = "\n".join(parts)
    truncated = len(full) > MAX_CHARS
    return {
        "text": full[:MAX_CHARS],
        "char_count": len(full),
        "has_tables": has_tables,
        "truncated": truncated,
    }


def extract_xlsx(file_path: str) -> dict:
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    sheets: list[dict] = []
    all_lines: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            str_row = [str(cell) if cell is not None else "" for cell in row]
            if any(c.strip() for c in str_row):
                rows.append(str_row)
                all_lines.append(" | ".join(str_row))
        sheets.append({"name": sheet_name, "rows": rows})
    wb.close()

    full_text = "\n".join(all_lines)
    truncated = len(full_text) > MAX_CHARS
    return {
        "sheets": sheets,
        "text_repr": full_text[:MAX_CHARS],
        "char_count": len(full_text),
        "truncated": truncated,
    }


def extract_text_file(file_path: str) -> dict:
    text = pathlib.Path(file_path).read_text(encoding="utf-8", errors="replace")
    truncated = len(text) > MAX_CHARS
    return {
        "text": text[:MAX_CHARS],
        "char_count": len(text),
        "truncated": truncated,
    }


def extract_image(file_path: str) -> dict:
    """Extract text description from an image using Claude Vision."""
    ext = pathlib.Path(file_path).suffix.lower()
    if ext not in IMAGE_EXTENSIONS:
        raise UnsupportedFileType(f"Unsupported image type: {ext!r}")

    media_type_map = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
    }
    media_type = media_type_map[ext]
    b64 = base64.standard_b64encode(pathlib.Path(file_path).read_bytes()).decode("utf-8")

    client = _get_anthropic_client()
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        messages=[{
            "role": "user",
            "content": [
                {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
                {"type": "text", "text": _VISION_PROMPT},
            ],
        }],
    )
    text = response.content[0].text

    # Use first heading as title, or fall back to filename stem
    title = pathlib.Path(file_path).stem
    for line in text.splitlines():
        stripped = line.lstrip("#").strip()
        if stripped:
            title = stripped
            break

    truncated = len(text) > MAX_CHARS
    return {
        "text": text[:MAX_CHARS],
        "title": title,
        "char_count": len(text),
        "truncated": truncated,
    }
